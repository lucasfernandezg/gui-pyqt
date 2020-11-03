import docx
from docx.oxml.xmlchemy import OxmlElement
from docx.oxml.shared import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()
sec_pr = doc.sections[0]._sectPr # get the section properties el
# create new borders el
pg_borders = OxmlElement('w:pgBorders')
# specifies how the relative positioning of the borders should be calculated
pg_borders.set(qn('w:offsetFrom'), 'page')
for border_name in ('top', 'left', 'bottom', 'right',): # set all borders
    border_el = OxmlElement(f'w:{border_name}')
    border_el.set(qn('w:val'), 'single') # a single line
    border_el.set(qn('w:sz'), '4') # for meaning of  remaining attrs please look docs
    border_el.set(qn('w:space'), '24')
    border_el.set(qn('w:color'), 'auto')
    pg_borders.append(border_el) # register single border to border el
sec_pr.append(pg_borders) # apply border changes to section

name ="Lucas"
h1 = doc.add_heading('Índice de reducción sonora de acuerdo con la Norma ISO 16283-1\nMedidas in situ del aislamiento a ruido aéreo entre recintos', 1)
p1 = doc.add_paragraph("____________________________________________________")
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
p.add_run('Name:'+name).alignment = "left"
p.add_run('Date').alignment = "right"



doc.save('border_test.docx')
